import docx
import csv
import re
import operator


def perform_computation(filename, file_stopwords, outputfile):
    doc = docx.Document(filename) #('/Users/hamza/Desktop/wordcount.docx')



    file = open(file_stopwords, 'r')
    stop_words = file.readlines()

    for i,j in enumerate(stop_words):
        stop_words[i] = stop_words[i].strip()

    x = 0
    y = 0
    birthdays = {}
    #declare a dictionary and then
    while(x < len(doc.paragraphs) - 1 and y < 101): #replace with len - 1
        print(y)
        string = ''
        number = ''
        for i in range(0, len(doc.paragraphs[x].text)):
            if(doc.paragraphs[x].text[i].isalpha() or doc.paragraphs[x].text[i] == "'"):
               string += doc.paragraphs[x].text[i]
            elif(doc.paragraphs[x].text[i] == ' ' or doc.paragraphs[x].text[i] == "’"):
                continue
            else:
                number += doc.paragraphs[x].text[i]

        string = string.strip()
        number = number.strip()

        number = re.sub('\s+', "'", number)
        number = int(number)

        if(string not in birthdays and len(string) > 4 and string not in stop_words):
            birthdays[string] = number
            y+=1
        elif(string in birthdays and len(string) > 4 and string not in stop_words):
            s = birthdays.get(string)
            s = int(s)
            number = int(number)
            total = s + number
            birthdays[string] = total
        x+=1


    sorted_d = dict(sorted(birthdays.items(), key=operator.itemgetter(1),reverse=True))
    counter = 0

    with open(outputfile, 'w', newline = '') as csv_file:
        writer = csv.writer(csv_file)
        for key, value in sorted_d.items():
          writer.writerow([key, value])
          counter = counter + 1
          if(counter == 100):
            break







